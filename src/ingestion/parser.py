"""Legal Document Parser (§3.1)
================================
Extracts raw text blocks from PDF (via PyMuPDF), DOCX (via python-docx),
and Markdown (from OCR), preserving formatting metadata for downstream
hierarchy reconstruction.

Output block schema:
    {
        "text":        str,
        "is_bold":     bool,
        "font_size":   float,
        "page_number": int,   # 1-indexed
    }

Font calibration derives a body-text baseline from real document statistics
so the ClauseChunker can distinguish headings from body text by size.
"""

from __future__ import annotations

from collections import Counter
from pathlib import Path
from typing import Any, Dict, List

import fitz  # PyMuPDF


class LegalDocumentParser:
    """Extracts formatted text blocks from legal documents.

    Supported formats: PDF, DOCX, Markdown (OCR output).
    """

    def __init__(self) -> None:
        self.body_font_size: float = 12.0

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def parse(self, filepath: str | Path) -> List[Dict[str, Any]]:
        """Route to the correct parser based on file extension."""
        path = Path(filepath)
        suffix = path.suffix.lower()
        parsers: dict[str, Any] = {
            ".pdf": self.parse_pdf,
            ".docx": self.parse_docx,
            ".md": self.parse_markdown,
        }
        fn = parsers.get(suffix)
        if fn is None:
            raise ValueError(
                f"Unsupported file type: {suffix} (supported: .pdf, .docx, .md)"
            )
        return fn(str(path))

    def parse_pdf(self, filepath: str) -> List[Dict[str, Any]]:
        """Extract blocks from a PDF with font calibration."""
        doc = fitz.open(filepath)
        self.body_font_size = self._calibrate_font_baseline(doc)

        extracted: List[Dict[str, Any]] = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            blocks = page.get_text("dict").get("blocks", [])
            for block in blocks:
                if "lines" not in block:
                    continue
                for line in block["lines"]:
                    for span in line["spans"]:
                        text = span["text"].strip()
                        if not text:
                            continue
                        extracted.append(
                            {
                                "text": text,
                                "is_bold": bool(span["flags"] & 2),
                                "font_size": span["size"],
                                "page_number": page_num + 1,  # 1-indexed
                            }
                        )
        doc.close()
        return extracted

    def parse_docx(self, filepath: str) -> List[Dict[str, Any]]:
        """Extract blocks from a DOCX using python-docx.

        Page number is approximated (30 paragraphs per page).
        Font size is calibrated from the modal run size of non-heading paragraphs.
        """
        from docx import Document  # local import avoids hard dep

        doc = Document(filepath)
        extracted: List[Dict[str, Any]] = []

        # Calibrate body font size from non-heading runs
        run_sizes: List[float] = []
        for para in doc.paragraphs:
            style_name = (para.style.name or "").lower()
            if not style_name.startswith("heading"):
                for run in para.runs:
                    if run.font.size:
                        run_sizes.append(round(run.font.size.pt, 1))
        self.body_font_size = (
            Counter(run_sizes).most_common(1)[0][0] if run_sizes else 12.0
        )

        PARAGRAPHS_PER_PAGE = 30
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            style_name = (para.style.name or "").lower()
            page_number = (idx // PARAGRAPHS_PER_PAGE) + 1

            # Heading-style paragraphs → bold + fixed size tiers
            heading_map: Dict[str, float] = {
                "heading 1": 16.0,
                "heading 2": 14.0,
                "heading 3": 13.0,
                "heading 4": 12.5,
            }
            if style_name in heading_map:
                is_bold = True
                font_size = heading_map[style_name]
            else:
                is_bold = any(run.bold for run in para.runs if run.bold is not None)
                sizes = [
                    run.font.size.pt for run in para.runs if run.font.size is not None
                ]
                font_size = max(sizes) if sizes else self.body_font_size

            extracted.append(
                {
                    "text": text,
                    "is_bold": is_bold,
                    "font_size": font_size,
                    "page_number": page_number,
                }
            )

        return extracted

    def parse_markdown(
        self, filepath: str | List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """Convert Markdown (from OCR/Gemini) into standard blocks.

        Accepts either a filepath to a .md file or a list of
        page dicts: [{"page": 1, "text": "## Article 1\n..."}, ...].
        """
        if isinstance(filepath, str):
            pages_data: List[Dict[str, Any]] = []
            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()
            # Treat the whole file as a single page if no explicit page splits
            pages_data = [{"page": 1, "text": content}]
        else:
            pages_data = filepath

        extracted: List[Dict[str, Any]] = []
        for page in pages_data:
            page_num = page.get("page", 1)
            for line in page.get("text", "").split("\n"):
                line = line.strip()
                if not line:
                    continue

                # Detect markdown headers
                if line.startswith("#"):
                    level = len(line) - len(line.lstrip("#"))
                    clean = line.lstrip("#").strip()
                    font_size = {1: 18.0, 2: 16.0, 3: 14.0}.get(level, 13.0)
                    extracted.append(
                        {
                            "text": clean,
                            "is_bold": True,
                            "font_size": font_size,
                            "page_number": page_num,
                        }
                    )
                else:
                    extracted.append(
                        {
                            "text": line,
                            "is_bold": False,
                            "font_size": self.body_font_size,
                            "page_number": page_num,
                        }
                    )
        return extracted

    # ------------------------------------------------------------------
    # Font calibration
    # ------------------------------------------------------------------

    def _calibrate_font_baseline(
        self, doc: fitz.Document, sample_pages: int = 5
    ) -> float:
        """Sample the first N pages and return the most common font size."""
        font_sizes: List[float] = []
        for page_num in range(min(sample_pages, len(doc))):
            page = doc[page_num]
            for block in page.get_text("dict").get("blocks", []):
                if "lines" not in block:
                    continue
                for line in block["lines"]:
                    for span in line["spans"]:
                        text = span["text"].strip()
                        if len(text) > 5:
                            font_sizes.append(round(span["size"], 1))
        if not font_sizes:
            return 12.0
        return Counter(font_sizes).most_common(1)[0][0]
